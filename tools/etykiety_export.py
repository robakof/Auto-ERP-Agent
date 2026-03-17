"""
etykiety_export.py — Generuje etykiety wysyłkowe Word (6x9cm) dla klienta z ERP.

Pobiera dane przez etykiety_10_oferty.sql, wypełnia szablon Word komórka po komórce.
Jeden plik = jeden klient, tyle stron ile potrzeba.

CLI:
    python tools/etykiety_export.py --klient-gid 9402 --output output/etykiety_AUCHAN.docx

Opcje:
    --klient-gid INT    GIDNumer grupy klienta (z etykiety_grupy.sql)
    --output PATH       Ścieżka do pliku wynikowego .docx
    --template PATH     Domyślnie: "Etykiety do wypełnienia.docx"
    --cols INT          Etykiet w wierszu (domyślnie: 4, jak w szablonie)
"""

import argparse
import copy
import re
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent))

from tools.lib.output import print_json
from tools.lib.sql_client import SqlClient

SQL_PATH = Path("solutions/jas/etykiety_10_oferty.sql")
TEMPLATE_PATH = Path("Etykiety do wypełnienia.docx")
DEFAULT_COLS = 4


# ---------------------------------------------------------------------------
# SQL
# ---------------------------------------------------------------------------

def _prepare_sql(klient_gid: int) -> str:
    """Usuwa DECLARE @klient_gid i podstawia wartość bezpośrednio w SQL."""
    sql = SQL_PATH.read_text(encoding="utf-8")
    sql = re.sub(
        r"DECLARE\s+@klient_gid\s+INT\s*=\s*\d+\s*;?\s*\n?",
        "",
        sql,
        flags=re.IGNORECASE,
    )
    sql = sql.replace("@klient_gid", str(int(klient_gid)))
    return sql.strip()


def _query_products(klient_gid: int) -> list[dict]:
    sql = _prepare_sql(klient_gid)
    result = SqlClient().execute(sql, inject_top=None)
    if not result["ok"]:
        raise RuntimeError(result["error"]["message"])
    cols = result["columns"]
    return [dict(zip(cols, row)) for row in result["rows"]]


# ---------------------------------------------------------------------------
# Wypełnianie komórki
# ---------------------------------------------------------------------------

def _v(value, suffix: str = "") -> str:
    """Zwraca wartość jako string z sufiksem, lub pusty string dla None/''."""
    if value is None:
        return ""
    s = str(value).strip()
    return (s + suffix) if s else ""


def _fill_cell(cell, product: dict) -> None:
    """
    Wypełnia komórkę szablonu danymi produktu.

    Struktura komórki (17 paragrafów, indeksy stałe dla szablonu):
      [2]  EAN:        [3]  → wartość EAN
      [5]  NAZWA:      [6]  → wartość nazwy
      [8]  Wysokość: {val} cm
      [9]  Czas palenia: {val} h
      [10] Gramatura: {val} g
      [12] COLI: {val} szt.
      [13] PALETA: {val} szt.
    """
    paras = cell.paragraphs

    # EAN — para[3]: dodaj run z wartością (para jest pusta w szablonie)
    _set_para_value(paras[3], _v(product.get("ean")))

    # NAZWA — para[6]
    _set_para_value(paras[6], _v(product.get("nazwa")))

    # Wysokość — para[8]: jeden run "Wysokość: "
    val = _v(product.get("wysokosc_cm"), " cm")
    _set_run_text(paras[8], 0, f"Wysokość: {val}")

    # Czas palenia — para[9]: dwa runy "Czas palenia" + ":"
    val = _v(product.get("czas_palenia_h"), " h")
    _set_run_text(paras[9], 0, f"Czas palenia: {val}")
    _set_run_text(paras[9], 1, "")

    # Gramatura — para[10]: dwa runy "Gramatura" + ":"
    val = _v(product.get("gramatura_g"), " g")
    _set_run_text(paras[10], 0, f"Gramatura: {val}")
    _set_run_text(paras[10], 1, "")

    # COLI — para[12]: jeden run "COLI:  "
    val = _v(product.get("coli_szt"))
    _set_run_text(paras[12], 0, f"COLI: {val} szt." if val else "COLI: ")

    # PALETA — para[13]: jeden run "PALETA: "
    val = _v(product.get("paleta_szt"))
    _set_run_text(paras[13], 0, f"PALETA: {val} szt." if val else "PALETA: ")


def _set_para_value(para, text: str) -> None:
    """Czyści istniejące runy i dodaje nowy z podanym tekstem."""
    for run in para.runs:
        run.text = ""
    if text:
        para.add_run(text)


def _set_run_text(para, index: int, text: str) -> None:
    """Ustawia tekst runu o podanym indeksie (jeśli istnieje)."""
    if index < len(para.runs):
        para.runs[index].text = text


# ---------------------------------------------------------------------------
# Generowanie dokumentu
# ---------------------------------------------------------------------------

def generate(
    products: list[dict],
    output_path: Path,
    template_path: Path = TEMPLATE_PATH,
    cols: int = DEFAULT_COLS,
) -> None:
    """
    Generuje plik Word z etykietami.

    Używa tabeli 0 z szablonu (siatka pustych etykiet, 4 kolumny).
    Wiersze dodawane dynamicznie — Word paginuje automatycznie.
    Pozostałe tabele z szablonu są usuwane.
    """
    doc = Document(str(template_path))
    body = doc.element.body

    table = doc.tables[0]
    tbl_el = table._tbl

    # Usuń wszystko po tabeli 0 (pozostałe tabele, akapity przykładów)
    # Zachowaj jeden akapit na końcu (Word tego wymaga)
    after = False
    to_remove = []
    for child in list(body):
        if child is tbl_el:
            after = True
            continue
        if after:
            to_remove.append(child)

    # Ostatni element body musi być akapitem — jeśli jest, zostaw go
    if to_remove and to_remove[-1].tag.endswith("}p"):
        to_remove = to_remove[:-1]
    for el in to_remove:
        body.remove(el)

    # Zachowaj wzorzec wiersza z szablonu, następnie usuń wszystkie wiersze
    template_row_xml = copy.deepcopy(table.rows[0]._tr)
    for row in list(table.rows):
        table._tbl.remove(row._tr)

    # Zbuduj tabelę: tyle wierszy ile potrzeba
    n = len(products)
    needed_rows = max(1, -(-n // cols))  # ceil(n / cols)

    for ri in range(needed_rows):
        new_row = copy.deepcopy(template_row_xml)
        table._tbl.append(new_row)

        row_obj = table.rows[ri]
        for ci in range(cols):
            idx = ri * cols + ci
            if idx < n:
                _fill_cell(row_obj.cells[ci], products[idx])
            # else: komórka zostaje pusta (format szablonu zachowany)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description="Generuj etykiety Word dla klienta ERP.")
    parser.add_argument("--klient-gid", type=int, required=True,
                        help="GIDNumer grupy klienta (z etykiety_grupy.sql)")
    parser.add_argument("--output", required=True,
                        help="Ścieżka do pliku wynikowego .docx")
    parser.add_argument("--template", default=str(TEMPLATE_PATH),
                        help=f"Szablon Word (domyślnie: {TEMPLATE_PATH})")
    parser.add_argument("--cols", type=int, default=DEFAULT_COLS,
                        help=f"Etykiet w wierszu (domyślnie: {DEFAULT_COLS})")
    args = parser.parse_args()

    output_path = Path(args.output)
    template_path = Path(args.template)

    if not template_path.exists():
        print_json({
            "ok": False, "data": None,
            "error": {"type": "TEMPLATE_NOT_FOUND", "message": str(template_path)},
            "meta": {},
        })
        return

    try:
        products = _query_products(args.klient_gid)
    except RuntimeError as e:
        print_json({
            "ok": False, "data": None,
            "error": {"type": "SQL_ERROR", "message": str(e)},
            "meta": {},
        })
        return

    if not products:
        print_json({
            "ok": False, "data": None,
            "error": {"type": "NO_DATA", "message": f"Brak produktów dla klient_gid={args.klient_gid}"},
            "meta": {},
        })
        return

    generate(products, output_path, template_path, args.cols)

    print_json({
        "ok": True,
        "data": {
            "row_count": len(products),
            "output_path": str(output_path.resolve()),
            "cols": args.cols,
            "rows_generated": -(-len(products) // args.cols),
        },
        "error": None,
        "meta": {},
    })


if __name__ == "__main__":
    main()
